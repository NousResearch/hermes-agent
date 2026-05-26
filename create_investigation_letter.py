from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# 设置页面边距
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

# 标题
title = doc.add_paragraph()
title_run = title.add_run('律师调查令申请书')
title_run.font.size = Pt(22)
title_run.font.bold = True
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 申请人信息
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('申请人：').bold = True
p.add_run('______律师事务所，地址：______')

p = doc.add_paragraph()
p.add_run('负责人：').bold = True
p.add_run('______，联系电话：______')

doc.add_paragraph()

# 申请事项
p = doc.add_paragraph()
p.add_run('申请事项：').bold = True
p.add_run('请求贵院开具律师调查令')

doc.add_paragraph()

# 请求依据
p = doc.add_paragraph()
p.add_run('请求依据：').bold = True

p = doc.add_paragraph()
p.add_run('申请人系______（原告/被告/第三人）与______（对方当事人）______纠纷一案中______（原告/被告/第三人）的委托诉讼代理人，执业证号：______。')

doc.add_paragraph()

# 调查事项
p = doc.add_paragraph()
p.add_run('调查事项：').bold = True

doc.add_paragraph('1. ______（被调查单位名称）')
doc.add_paragraph('   调查内容：______')
doc.add_paragraph('   与本案的关联性：______')

doc.add_paragraph('2. ______（其他调查事项）')

doc.add_paragraph()

# 申请理由
p = doc.add_paragraph()
p.add_run('申请理由：').bold = True

p = doc.add_paragraph()
p.add_run('上述证据材料由______（单位/个人）持有，申请人因______（客观原因）无法自行收集。该证据材料对查清本案事实______（明确证明目的）具有重要作用，属于《中华人民共和国民事诉讼法》第六十七条规定 的"当事人及其诉讼代理人因客观原因不能自行收集的其他证据"。')

doc.add_paragraph()

# 证据线索
p = doc.add_paragraph()
p.add_run('证据线索：').bold = True

doc.add_paragraph('• 被调查单位：______，地址：______')
doc.add_paragraph('• 联系人/联系方式：______')
doc.add_paragraph('• 已知线索：______')

doc.add_paragraph()

# 承诺事项
p = doc.add_paragraph()
p.add_run('承诺事项：').bold = True

p = doc.add_paragraph()
p.add_run('申请人承诺将妥善保管调查令及调取的材料，仅用于本案诉讼目的，不挪作他用，并在开庭审理前将调取证据如实提交贵院。')

doc.add_paragraph()
doc.add_paragraph()

# 致送法院
p = doc.add_paragraph('此致')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

p = doc.add_paragraph('______人民法院')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# 落款
p = doc.add_paragraph()
p.add_run('申请人：______律师事务所\n')
p.add_run('代理人：______（签名）\n')
p.add_run('______年______月______日')
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

path = '/Users/appleoppa/Desktop/律师调查令申请书.docx'
doc.save(path)
print('已保存：' + path)
